from __future__ import annotations

import argparse
import difflib
import re
from html import escape
from pathlib import Path

from docx import Document


def read_docx_paragraphs(path: Path) -> list[str]:
    """docxから段落単位のリストを取得する（既存read_docxと同じ抽出ロジック）。"""
    doc = Document(path)
    paragraphs: list[str] = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            paragraphs.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" / ".join(cells))

    return paragraphs


def extract_revised_section(md_text: str) -> str:
    """Claude出力mdから「## 3. 修正版全文」セクションのみを抜き出す。
    見つからない場合は全文を対象にする（フォーマット崩れへの保険）。
    """
    lines = md_text.splitlines()
    start: int | None = None
    end = len(lines)

    for i, line in enumerate(lines):
        if re.match(r"^##\s*3\.", line):
            start = i + 1
        elif start is not None and re.match(r"^##\s*4\.", line):
            end = i
            break

    if start is None:
        return md_text

    return "\n".join(lines[start:end]).strip()


def to_paragraphs(text: str) -> list[str]:
    return [p.strip() for p in text.split("\n\n") if p.strip()]


def inline_highlight(before_text: str, after_text: str) -> str:
    """1組の段落（before/after）を文字単位で比較し、
    afterをベースに『追加=緑』『変更=赤』のみをHTMLでマークする。
    afterに存在しない部分（削除）は出力しない。
    """
    sm = difflib.SequenceMatcher(None, before_text, after_text, autojunk=False)
    out: list[str] = []

    for tag, _i1, _i2, j1, j2 in sm.get_opcodes():
        chunk = after_text[j1:j2]
        if not chunk:
            continue
        if tag == "equal":
            out.append(escape(chunk))
        elif tag == "insert":
            out.append(f'<span class="added">{escape(chunk)}</span>')
        elif tag == "replace":
            out.append(f'<span class="changed">{escape(chunk)}</span>')
        # tag == "delete" は after に存在しないため何もしない

    return "".join(out)


def build_highlighted_html(before_paragraphs: list[str], after_paragraphs: list[str]) -> str:
    """段落単位でbefore/afterを対応づけ、afterの文章のみを
    ハイライト付きHTML段落として組み立てる。"""
    sm = difflib.SequenceMatcher(None, before_paragraphs, after_paragraphs, autojunk=False)
    blocks: list[str] = []

    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            for para in after_paragraphs[j1:j2]:
                blocks.append(f"<p>{escape(para)}</p>")

        elif tag == "insert":
            # before側に対応がない＝完全新規の段落
            for para in after_paragraphs[j1:j2]:
                blocks.append(f'<p><span class="added">{escape(para)}</span></p>')

        elif tag == "delete":
            # afterに存在しない段落なので表示しない
            continue

        elif tag == "replace":
            before_block = before_paragraphs[i1:i2]
            after_block = after_paragraphs[j1:j2]
            paired = min(len(before_block), len(after_block))

            # 対応する段落どうしは文字単位で差分ハイライト
            for k in range(paired):
                html_text = inline_highlight(before_block[k], after_block[k])
                blocks.append(f"<p>{html_text}</p>")

            # afterの段落数がbeforeより多い場合、残りは完全新規として緑表示
            for extra_para in after_block[paired:]:
                blocks.append(f'<p><span class="added">{escape(extra_para)}</span></p>')

    return "\n".join(blocks)


CUSTOM_CSS = """
<style>
body {
  font-family: -apple-system, "Hiragino Sans", "Yu Gothic", sans-serif;
  line-height: 1.9;
  padding: 2em;
  max-width: 760px;
  margin: auto;
  color: #222;
  font-size: 16px;
}
h2 { margin-bottom: 0.2em; }
p.meta { color: #999; font-size: 0.85em; margin-top: 0; }
.legend { font-size: 0.85em; color: #666; margin-bottom: 1.5em; }
.legend .added { padding: 0 0.3em; }
.legend .changed { padding: 0 0.3em; }
.added { color: #1a7f37; font-weight: 600; }
.changed { color: #c1121f; font-weight: 600; }
p { margin: 0 0 1em 0; }
</style>
"""


def main() -> None:
    parser = argparse.ArgumentParser(
        description="校正前(.docx)と校正後(Claude出力.md)を比較し、"
        "校正後の文章のみを対象に「追加=緑字」「変更=赤字」でハイライトしたHTMLを生成します。"
    )
    parser.add_argument("--before", required=True, help="校正前の.docxファイル")
    parser.add_argument("--after", required=True, help="校正後の.mdファイル（Claude/GPT出力）")
    parser.add_argument("--output", required=True, help="出力する.htmlファイル")

    args = parser.parse_args()

    before_path = Path(args.before).expanduser().resolve()
    after_path = Path(args.after).expanduser().resolve()
    output_path = Path(args.output).expanduser().resolve()

    if not before_path.exists():
        raise FileNotFoundError(f"File not found: {before_path}")
    if not after_path.exists():
        raise FileNotFoundError(f"File not found: {after_path}")

    before_paragraphs = read_docx_paragraphs(before_path)

    after_full_text = after_path.read_text(encoding="utf-8")
    revised_text = extract_revised_section(after_full_text)
    after_paragraphs = to_paragraphs(revised_text)

    body_html = build_highlighted_html(before_paragraphs, after_paragraphs)

    html = (
        "<!DOCTYPE html><html><head><meta charset='utf-8'>"
        f"{CUSTOM_CSS}"
        "</head><body>"
        "<h2>校正後の文章（ハイライト付き）</h2>"
        f"<p class='meta'>before: {before_path.name} / after: {after_path.name}</p>"
        '<p class="legend">'
        '<span class="added">緑字＝新規追加</span>'
        '<span class="changed">赤字＝既存部分の修正</span>'
        "（無色＝変更なし。本機能は補助的な可視化のため、正式な修正点は本文のレビューを優先してください）"
        "</p>"
        f"{body_html}"
        "</body></html>"
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(html, encoding="utf-8")

    print(f"[OK] ハイライト生成: {output_path}")


if __name__ == "__main__":
    main()