from __future__ import annotations

import argparse
from pathlib import Path
from docx import Document


def read_md(path: Path | None) -> str:
    if path is None:
        return ""
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")
    return path.read_text(encoding="utf-8")


def read_docx(path: Path) -> str:
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    doc = Document(path)
    paragraphs: list[str] = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            paragraphs.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" / ".join(cells))

    return "\n\n".join(paragraphs)


def build_prompt(
    article_text: str,
    policy_text: str,
    author_text: str = "",
    reader_text: str = "",
) -> str:
    parts: list[str] = []

    parts.append("# GPT校正依頼")
    parts.append("")
    parts.append("以下の校正指針に従って、note記事を校正してください。")
    parts.append("目的は、文章を一般的に綺麗にすることではなく、発信者の思想・温度感・科学的誠実さを保ちながら、読者に届く文章へ整えることです。")
    parts.append("")

    if author_text:
        parts.append("---")
        parts.append("")
        parts.append("# 発信者プロフィール")
        parts.append("")
        parts.append(author_text)

    if reader_text:
        parts.append("---")
        parts.append("")
        parts.append("# 読者ペルソナ")
        parts.append("")
        parts.append(reader_text)

    parts.append("---")
    parts.append("")
    parts.append("# 校正指針")
    parts.append("")
    parts.append(policy_text)

    parts.append("---")
    parts.append("")
    parts.append("# 新規記事本文")
    parts.append("")
    parts.append(article_text)

    parts.append("---")
    parts.append("")
    parts.append("# 出力形式")
    parts.append("")
    parts.append("以下の形式で出力してください。")
    parts.append("")
    parts.append("## 1. 全体レビュー")
    parts.append("- 記事全体の良い点")
    parts.append("- 改善すべき点")
    parts.append("- 発信者らしさが保たれているか")
    parts.append("")
    parts.append("## 2. 校正方針")
    parts.append("- どの観点を優先して直したか")
    parts.append("- どの表現をあえて残したか")
    parts.append("- どの表現を削った、または弱めたか")
    parts.append("")
    parts.append("## 3. 修正版全文")
    parts.append("- noteへ貼り付けられるMarkdown形式")
    parts.append("- 見出し、段落、余白を整える")
    parts.append("- 元の思想・熱量を壊さない")
    parts.append("")
    parts.append("## 4. 人間確認ポイント")
    parts.append("- 投稿前に本人が確認すべき点")
    parts.append("- 科学的事実確認が必要な箇所")
    parts.append("- 表現の強さを最終判断すべき箇所")

    return "\n".join(parts)


def main() -> None:
    parser = argparse.ArgumentParser(
        description="新規note記事のGPT校正用プロンプトを生成します。"
    )

    parser.add_argument("--article", required=True, help="新規記事のdocxファイル")
    parser.add_argument("--policy", required=True, help="editorial_policy.md")
    parser.add_argument("--output", required=True, help="出力するprompt.md")

    parser.add_argument("--author", default=None, help="author_persona.md 任意")
    parser.add_argument("--reader", default=None, help="target_reader.md 任意")

    args = parser.parse_args()

    article_path = Path(args.article).expanduser().resolve()
    policy_path = Path(args.policy).expanduser().resolve()
    output_path = Path(args.output).expanduser().resolve()

    author_path = Path(args.author).expanduser().resolve() if args.author else None
    reader_path = Path(args.reader).expanduser().resolve() if args.reader else None

    article_text = read_docx(article_path)
    policy_text = read_md(policy_path)
    author_text = read_md(author_path)
    reader_text = read_md(reader_path)

    prompt = build_prompt(
        article_text=article_text,
        policy_text=policy_text,
        author_text=author_text,
        reader_text=reader_text,
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(prompt, encoding="utf-8")

    print(f"[OK] Generated: {output_path}")


if __name__ == "__main__":
    main()